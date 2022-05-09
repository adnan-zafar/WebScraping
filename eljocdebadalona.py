import requests
from bs4 import BeautifulSoup
import pandas as pd
import time
from docx import Document
from docx.shared import Inches
import json
import concurrent.futures
from datetime import datetime

cookies = {
    'PHPSESSID': 'jnerl4c6tc82gqgrc8582lq1b7',
    '__utmc': '164870891',
    '__utmz': '164870891.1650983811.1.1.utmcsr=(direct)^|utmccn=(direct)^|utmcmd=(none)',
    'arp_scroll_position': '0',
    '__utma': '164870891.698289709.1650983811.1650983811.1650990724.2',
    '__utmt': '1',
    '__utmb': '164870891.1.10.1650990724',
}

headers = {
    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
    'Accept-Language': 'en-PK,en-US;q=0.9,en;q=0.8',
    'Cache-Control': 'max-age=0',
    'Connection': 'keep-alive',
    'Referer': 'http://www.eljocdebadalona.cat/preguntes/historiclist',
    'Upgrade-Insecure-Requests': '1',
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Safari/537.36',
}

base_url = 'http://www.eljocdebadalona.cat'
list_of_main_page_urls = ['/preguntes/historicview/12', '/preguntes/historicview/11','/preguntes/historicview/10','/preguntes/historicview/9', '/preguntes/historicview/8','/preguntes/historicview/7','/preguntes/historicview/6', '/preguntes/historicview/5', '/preguntes/historicview/1', '/preguntes/historicview/2', '/preguntes/historicview/3', '/preguntes/historicview/4']
list_of_dict = []
skip_url = 'http://www.eljocdebadalona.cat/preguntes/2015-05-08/a-la-nostra-ciutat-hi-ha-el-passatge-d-en-sampere-a-qui'

def pagescraper(i):
    single_page_url = base_url+i
    print(single_page_url)
    if single_page_url == skip_url:
        print("Found URL: ", i)
    else:
        dictionary = {}
        res = requests.get(single_page_url,headers=headers, cookies=cookies, verify=False)
        soup = BeautifulSoup(res.content, 'lxml')
        print(len(soup.select('div#contingut_txt>p')))
        dictionary['Date']       = soup.select_one('div#contingut_txt>h4').text
        dictionary['Option A']   = soup.select('div#contingut_txt>ol.anteriors_respostes>li')[0].text
        dictionary['Option B']   = soup.select('div#contingut_txt>ol.anteriors_respostes>li')[1].text
        dictionary['Option C']   = soup.select('div#contingut_txt>ol.anteriors_respostes>li')[2].text
        if len(soup.select('div#contingut_txt>p')) == 7:
            dictionary['Answer']    = soup.select('div#contingut_txt>p')[3].text
            if soup.select('div#contingut_txt>p')[0].text == '':
                dictionary['Question']   = soup.select('div#contingut_txt>p')[1].text 
            else:                            
                question_one   = soup.select('div#contingut_txt>p')[0].text
                question_two   = soup.select('div#contingut_txt>p')[1].text
                dictionary['Question']   = question_one+'\n'+question_two
        elif len(soup.select('div#contingut_txt>p')) == 9:
            question_one   = soup.select('div#contingut_txt>p')[1].text
            question_two   = soup.select('div#contingut_txt>p')[2].text
            dictionary['Question']   = question_one+'\n'+question_two
            dictionary['Answer']    = soup.select('div#contingut_txt>p')[5].text
        elif len(soup.select('div#contingut_txt>p')) == 8:
            question_one   = soup.select('div#contingut_txt>p')[1].text
            question_two   = soup.select('div#contingut_txt>p')[2].text
            dictionary['Question']   = question_one+'\n'+question_two
            dictionary['Answer']    = soup.select('div#contingut_txt>p')[4].text
        else:                      
            dictionary['Question']   = soup.select('div#contingut_txt>p')[0].text            
            dictionary['Answer']    = soup.select('div#contingut_txt>p')[2].text
    list_of_dict.append(dictionary)
    print(dictionary)

def scraper(i):
    category_urls = list()
    main_page_url = base_url+i
    req_main_page = requests.get(main_page_url,headers=headers, cookies=cookies, verify=False)
    main_soup = BeautifulSoup(req_main_page.content,'lxml')
    last_url =[i.get('href') for i in main_soup.select('li.last>a')]
    last_page_num = last_url[0].split('=')[-1]
    print(last_page_num,main_page_url)
    for x in range(1,int(last_page_num)+1):
        inner_url = f'{main_page_url}?page={x}'
        main_res = requests.get(inner_url.strip(), headers=headers, cookies=cookies, verify=False)
        main_soup = BeautifulSoup(main_res.content, 'lxml')
        links = [i.get('href') for i in main_soup.select('p.pregunta-ant>a')]
        category_urls.extend(links)
    with concurrent.futures.ThreadPoolExecutor() as executor:
        executor.map(pagescraper, category_urls)
    


if __name__ == '__main__':
    print('\nDate & Time ::', datetime.now())  
    with concurrent.futures.ThreadPoolExecutor() as executor:
        executor.map(scraper, list_of_main_page_urls)

    for i in list_of_dict:
        document.add_paragraph(i["Date"])
        document.add_paragraph(i["Question"])
        document.add_paragraph(i["Option A"],style='List Number')
        document.add_paragraph(i["Option B"],style='List Number')
        document.add_paragraph(i["Option C"],style='List Number')
        document.add_paragraph(i["Answer"])

    document.save('eljocdebadalona.docx')
    print('Process Complete...')
    print('\nDate & Time ::', datetime.now())
        